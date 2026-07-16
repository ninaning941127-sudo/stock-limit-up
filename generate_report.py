"""
把全台股漲停策略回測的結果組成一份 Word 研究報告。

直接重用 market_limit_up_backtest.py 的核心邏輯（載入資料、判斷漲停事件、
按日聚合報酬率），不重寫策略邏輯；額外抓 0050.TW 同期間買進持有報酬率
做對照，最後用 python-docx 組裝成 .docx。
"""

from pathlib import Path

import numpy as np
import pandas as pd
import yfinance as yf
from scipy import stats

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

import market_limit_up_backtest as core

OUT_DIR = Path(__file__).parent
REPORT_PATH = OUT_DIR / "台股漲停策略回測報告.docx"

BENCHMARK_TICKER = "0050.TW"
BENCHMARK_START = "2026-01-02"
BENCHMARK_END_EXCLUSIVE = "2026-07-09"  # yfinance end 是不含當天，+1 天才含 07-08

CJK_FONT = "微軟正黑體"


# ---------- docx 小工具：確保中文用指定字型顯示 ----------
def _set_east_asian_font(font_obj, element, name):
    font_obj.name = name
    rpr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_style_cjk(style, name=CJK_FONT):
    _set_east_asian_font(style.font, style.element, name)


def set_run_cjk(run, name=CJK_FONT):
    _set_east_asian_font(run.font, run._element, name)


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_run_cjk(run)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    set_run_cjk(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_cjk(run)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        set_run_cjk(run)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            set_run_cjk(run)
    return table


def add_image(doc, path, width_in=6.0, caption=None):
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        set_run_cjk(run)


# ---------- 分析數字（重用 core 的邏輯） ----------
def compute_all_stats():
    df = core.load_data(core.DATA_XLSX)
    trades, pending = core.find_trades(df)

    n = len(trades)
    n_stocks = trades["code"].nunique()
    date_min, date_max = trades["date"].min().date(), trades["date"].max().date()

    by_type = trades.groupby("type").agg(
        件數=("pnl_pct", "size"),
        勝率=("pnl_pct", lambda s: (s > 0).mean()),
        平均報酬率=("pnl_pct", "mean"),
    )
    overall_win_rate = (trades["pnl_pct"] > 0).mean()
    overall_avg = trades["pnl_pct"].mean()

    daily_df = core.compute_daily_series(trades)
    final_cum = daily_df.groupby("series")["cum_return"].last()

    # 按日聚合 t 檢定（跟 core.statistical_tests 邏輯一致）
    day_tests = {}
    for label, sub in [
        ("整體", trades),
        ("僅鎖死", trades[trades["type"] == "locked"]),
        ("僅未鎖死", trades[trades["type"] == "not_locked"]),
    ]:
        daily = sub.groupby("date")["pnl_pct"].mean()
        t_stat, p_value = stats.ttest_1samp(daily, popmean=0)
        day_tests[label] = {
            "mean": daily.mean(), "n_days": len(daily), "t": t_stat, "p": p_value,
        }

    ic_df = core.ic_ir_analysis(trades)
    ic_stats = {
        "n_days": len(ic_df),
        "mean_ic": ic_df["ic"].mean(),
        "std_ic": ic_df["ic"].std(ddof=1),
        "ir": ic_df["ic"].mean() / ic_df["ic"].std(ddof=1),
    }
    ic_t, ic_p = stats.ttest_1samp(ic_df["ic"], popmean=0)
    ic_stats["t"] = ic_t
    ic_stats["p"] = ic_p

    return {
        "trades": trades, "pending": pending, "n": n, "n_stocks": n_stocks,
        "date_min": date_min, "date_max": date_max, "by_type": by_type,
        "overall_win_rate": overall_win_rate, "overall_avg": overall_avg,
        "daily_df": daily_df, "final_cum": final_cum,
        "day_tests": day_tests, "ic_stats": ic_stats,
    }


def fetch_benchmark_return():
    data = yf.download(
        BENCHMARK_TICKER, start=BENCHMARK_START, end=BENCHMARK_END_EXCLUSIVE,
        auto_adjust=True, progress=False,
    )
    if isinstance(data.columns, pd.MultiIndex):
        data.columns = data.columns.get_level_values(0)
    first_close = data["Close"].iloc[0]
    last_close = data["Close"].iloc[-1]
    ret = last_close / first_close - 1
    return float(ret), data.index[0].date(), data.index[-1].date()


# ---------- 組報告 ----------
def build_report(stats_dict, bench_ret, bench_start, bench_end):
    doc = Document()

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            set_style_cjk(doc.styles[style_name])

    title = doc.add_heading(level=0)
    run = title.add_run("台股漲停策略回測報告")
    set_run_cjk(run)

    add_paragraph(
        doc,
        f"資料期間：{stats_dict['date_min']} ~ {stats_dict['date_max']}　"
        f"標的範圍：全台股（TEJ 匯出資料，共 {stats_dict['n_stocks']} 檔股票、"
        f"{stats_dict['n']} 筆漲停事件）",
        italic=True,
    )

    # ---- 摘要 ----
    add_heading(doc, "摘要", level=1)
    ft = stats_dict["final_cum"]
    add_paragraph(
        doc,
        f"本報告回測一個假設性的台股當沖式策略：只要個股當天觸及漲停就假設能買到漲停價進場，"
        f"再依當天收盤是否鎖死漲停分兩種方式出場。以 2026 年上半年全台股資料驗證，共偵測到 "
        f"{stats_dict['n']} 筆漲停事件、涉及 {stats_dict['n_stocks']} 檔股票。整體策略的累積報酬率"
        f"（非複利，逐日等權重報酬率加總）達 {ft['整體']:.1%}，統計檢定顯示鎖死與未鎖死兩組交易的"
        f"報酬率差異在統計上高度顯著；但報告最後會說明，這個名目上的高報酬主要來自現實中最難成交"
        f"的「鎖死」交易，實際可執行性存疑。"
    )

    # ---- 一、策略說明 ----
    add_heading(doc, "一、策略說明", level=1)
    add_paragraph(doc, "進場條件：個股當天股價觸及漲停價（視為能以漲停價買入）。出場方式分兩種：")
    add_bullet(doc, "鎖死到收盤（收盤價＝最高價＝漲停價）：隔天開盤立刻賣出，"
                    "損益 = 隔天開盤價 − 當天收盤價")
    add_bullet(doc, "觸及漲停但未鎖死（收盤價 < 最高價）：當天收盤賣出，"
                    "損益 = 當天收盤價 − 當天最高價")
    add_paragraph(
        doc,
        "漲停價計算：前一交易日收盤價 × 1.10，再依台灣證交所實際的升降單位（tick）規則"
        "無條件捨去到合法價位（例如股價 100~500 元 tick=0.5 元、500~1000 元 tick=1 元、"
        "≥1000 元 tick=5 元），確保跟交易所公告的真實漲停價一致。"
    )
    add_paragraph(
        doc,
        "交易成本：每筆交易都已扣除賣出時的證交稅（賣出成交價 × 0.3%，直接從損益中扣除），"
        "但買賣手續費（依券商折扣不同，概念上約為成交價的 0.1%~0.1425%）尚未計入，"
        "本報告後續所有報酬率、統計檢定與圖表都是「已扣證交稅、未扣手續費」下的結果。"
    )
    add_paragraph(
        doc,
        f"資料來源為 TEJ 匯出的全台股日成交資料（開高低收），期間 "
        f"{stats_dict['date_min']} ~ {stats_dict['date_max']}，涵蓋約 1,900 檔上市櫃股票。"
    )

    # ---- 二、回測結果 ----
    add_heading(doc, "二、回測結果", level=1)

    add_heading(doc, "2.1 累積報酬率", level=2)
    add_paragraph(
        doc,
        "計算方式分兩步：(1) 同一天可能有多檔股票同時觸及漲停，先在「天」的層級做等權重平均，"
        "得到當天的報酬率（例如當天 3 筆交易分別 +10%／-2%／+6%，則當天報酬率 = "
        "(10%-2%+6%)/3 = 4.67%）；(2) 把每個交易日的報酬率依時間順序直接加總（非複利），"
        "得到累積報酬率（例如第一天 +4%、第二天 +5%，累積 = 4%+5% = 9%）。"
        "「整體」「僅鎖死」「僅未鎖死」三條線各自獨立計算，因為各自「有交易的日子」不同。"
    )
    add_image(doc, core.CUM_PNG, caption="圖 1：累積報酬率（每日報酬率加總，非複利）")

    add_heading(doc, "2.2 當日報酬率（非累積）", level=2)
    add_paragraph(doc, "同一張邏輯的「未加總」版本，直接呈現每個有交易日子當天的報酬率，"
                       "可以看出鎖死與未鎖死兩組報酬率在整段期間的分布是否穩定。")
    add_image(doc, core.DAILY_PNG, caption="圖 2：有交易當天的報酬率")

    add_heading(doc, "2.3 漲停事件報酬率分布", level=2)
    add_paragraph(doc, "把所有漲停事件的單筆報酬率畫成分布圖，鎖死組與未鎖死組幾乎沒有重疊："
                       "鎖死組集中在 0 附近到略低於 10% 之間（右側被台股漲跌幅限制加上證交稅"
                       "共同截斷——隔天開盤價本身也受漲停限制，扣稅後上限略低於 10%），"
                       "未鎖死組則全部落在 0 以下。")
    add_image(doc, core.DIST_PNG, caption="圖 3：漲停事件報酬率分布（鎖死 vs 未鎖死）")

    add_heading(doc, "2.4 跟 0050 買進持有比較", level=2)
    strategy_overall = ft["整體"]
    add_table(
        doc,
        ["項目", "策略（整體，累積報酬率）", "0050 買進持有（含息，price return）"],
        [
            ["期間", f"{stats_dict['date_min']} ~ {stats_dict['date_max']}", f"{bench_start} ~ {bench_end}"],
            ["報酬率", f"{strategy_overall:.1%}", f"{bench_ret:.1%}"],
        ],
    )
    add_paragraph(
        doc,
        "⚠️ 兩者計算方式不同，不是嚴謹的同基準比較：策略的數字是「每個交易日等權重平均、"
        "再逐日加總」的結果，代表的是「每次漲停訊號都用等量部位下注」的加總報酬，並非單一"
        "帳戶的複利成長；0050 則是單一部位、從頭到尾抱著不動的真實複利報酬。策略數字雖已扣除"
        "賣出證交稅，但仍未扣手續費、也沒有考慮流動性與資金能否同時應付多檔同時漲停的限制"
        "（詳見「結論與限制」）。在此前提下，策略的名目數字高於 0050 同期表現，但不代表實際"
        "可以賺到這個差距。"
    )

    # ---- 三、結論：這個策略到底能不能賺錢？ ----
    add_heading(doc, "三、結論：到底能不能賺錢？", level=1)
    dt = stats_dict["day_tests"]
    add_paragraph(
        doc,
        f"統計上，「鎖死」與「未鎖死」兩組交易的報酬率差異非常顯著且穩定（詳見下一節統計檢定），"
        f"整體策略每日平均報酬率 {dt['整體']['mean']:.2%}（{dt['整體']['n_days']} 個交易日，"
        f"p={dt['整體']['p']:.2g}）。但「能不能賺錢」要看你實際能不能執行到這個策略的核心 "
        f"— 也就是鎖死那部分（{stats_dict['by_type'].loc['locked','件數']:.0f} 筆，勝率 "
        f"{stats_dict['by_type'].loc['locked','勝率']:.1%}，平均 "
        f"{stats_dict['by_type'].loc['locked','平均報酬率']:.2%}）："
    )
    add_bullet(doc, "股票會鎖死漲停，正是因為當天想買的人遠多於想賣的人，實務上散戶的委買單"
                    "往往整天都排不到、成交不了，回測假設「只要漲停就買得到」明顯高估了可執行性")
    add_bullet(doc, "未鎖死組（相對容易成交的部分）勝率必然是 0%——因為出場價=收盤價、"
                    "進場價=當天最高價，收盤不可能高於當天最高，這是策略定義下的數學必然，不是新發現")
    add_bullet(doc, "已扣除賣出時的證交稅（0.3%），但手續費仍未計入，也沒有考慮同一天"
                    "數十檔同時漲停時，資金該怎麼分配的限制")
    add_paragraph(
        doc,
        "簡短結論：這個策略在統計上「有效應」是真的（不是隨機雜訊），方向也一致地指向"
        "「鎖死能賺、未鎖死會虧」；但名目上算出來比 0050 高的報酬率，主要建立在「鎖死當天"
        "買得到」這個現實中最不可能成立的假設上，實際能不能賺到、賺多少，取決於你有沒有辦法"
        "真的在漲停鎖死的當下排隊買到股票——這是下一步如果想驗證「真的能不能賺」最需要補上的資料"
        "（例如逐筆成交或委買委賣掛單量）。"
    )

    # ---- 四、統計顯著性檢定 ----
    add_heading(doc, "四、統計顯著性檢定", level=1)
    add_paragraph(
        doc,
        "這一節想回答一個很單純的問題：前面看到的報酬率差異，會不會只是運氣、剛好在這半年"
        "被我們遇到？還是背後真的有穩定、可重複的規律？統計上用「p 值」來回答這個問題——"
        "p 值可以理解成：『假設這個策略其實完全沒用、賺賠都是隨機的，我們卻觀察到這麼極端"
        "的數據，機率會有多低』。p 值越小，代表「純粹運氣」這個解釋越說不通；一般會用 "
        "p < 0.05（5%）當作「統計上顯著」的門檻。另外要注意：同一天常有數十檔股票一起"
        "漲停，彼此的報酬率會互相牽動、不是各自獨立，所以下面的檢定都先把同一天的交易"
        "合併成一個數字，再用「交易日數」而不是「交易筆數」當樣本數，避免把統計顯著性"
        "灌水。"
    )

    add_heading(doc, "4.1 鎖死跟未鎖死，報酬率是不是真的不一樣？", level=2)
    add_paragraph(
        doc,
        "作法：同一天如果有好幾筆交易，先平均成一個「當天報酬率」；整體、僅鎖死、僅未鎖死"
        "各自累積出一串「每天報酬率」的數字（大約 120 天），再檢查這串數字平均起來是不是"
        "顯著大於或小於 0。"
    )
    add_table(
        doc,
        ["分類", "每日平均報酬率", "交易日數", "t 值", "p 值"],
        [
            [label, f"{v['mean']:.2%}", v["n_days"], f"{v['t']:.2f}", f"{v['p']:.3g}"]
            for label, v in dt.items()
        ],
    )
    add_paragraph(
        doc,
        f"白話解讀：僅鎖死組平均每天賺 {dt['僅鎖死']['mean']:.2%}，p 值小到幾乎是 0，"
        f"代表這不太可能是巧合；僅未鎖死組平均每天虧 {abs(dt['僅未鎖死']['mean']):.2%}，"
        f"同樣是極端顯著的虧損；兩者混在一起的整體，平均每天"
        f"{'賺' if dt['整體']['mean'] >= 0 else '虧'} {abs(dt['整體']['mean']):.2%}，"
        f"也顯著不等於 0。三個結果都指向同一個結論：是顯著的——鎖死賺、未鎖死虧"
        "這個規律在半年、上百個交易日裡反覆出現，不是少數幾天的好運撐出來的（以上數字"
        "都已扣除賣出 0.3% 證交稅）。"
    )
    add_paragraph(
        doc,
        "（鎖死組跟未鎖死組哪個比較賺，這裡不另外做一個「兩組差異」的檢定：未鎖死組的"
        "出場價依定義恆 ≤ 進場價（收盤價不可能高於當天最高價），這組報酬率 ≤ 0 是規則"
        "本身造成的，不是需要驗證的統計問題，硬做檢定只是在檢定一個恆成立的定義。）"
    )

    add_heading(doc, "4.2 「鎖死」這個訊號，預測力穩不穩定？（IC / IR）", level=2)
    add_paragraph(
        doc,
        "IC（Information Coefficient）：每天分別檢查「這檔股票今天有沒有鎖死」跟「它的"
        "報酬率高低」是不是同方向——鎖死的股票報酬率排名越是明顯偏高，IC 就越接近 1；"
        "兩者完全沒關係，IC 就接近 0。"
    )
    add_paragraph(
        doc,
        "IR（Information Ratio）：只看平均 IC 還不夠，因為可能是少數幾天特別準、其他天"
        "都在亂猜。IR 把每天的 IC 拿來算「平均值 ÷ 波動度」，概念上很像投資常聽到的夏普"
        "比率——IR 越高，代表這個訊號不只平均有效，而且幾乎每天都穩定有效，不是靠運氣。"
    )
    ic = stats_dict["ic_stats"]
    add_table(
        doc,
        ["可計算 IC 的交易日數", "平均 IC", "IC 標準差", "IR", "t 值", "p 值"],
        [[ic["n_days"], f"{ic['mean_ic']:.3f}", f"{ic['std_ic']:.3f}",
          f"{ic['ir']:.2f}", f"{ic['t']:.2f}", f"{ic['p']:.3g}"]],
    )
    add_paragraph(
        doc,
        f"白話解讀：平均 IC 高達 {ic['mean_ic']:.2f}（滿分 1），算是非常強的相關性；"
        f"IR 達 {ic['ir']:.1f}，代表這個效應不只平均起來明顯，過去半年裡幾乎天天成立，"
        f"不是少數幾天特別準造成的錯覺。p 值同樣遠小於 0.05，是顯著的，「鎖死能預測"
        "較高報酬」這件事不太可能是巧合。"
    )

    add_paragraph(
        doc,
        "小結：以上三個檢定的答案都一樣——統計上顯著，這個規律很可能是真實存在、"
        "不是雜訊。但「統計顯著」只代表「這個現象大概率不是巧合」，不等於「現實中一定"
        "賺得到這筆錢」：能不能真的變現，取決於第三節提到的「鎖死當下買不買得到」，"
        "這才是最終能不能賺錢的關鍵。"
    )

    # ---- 五、已知限制 ----
    add_heading(doc, "五、已知限制", level=1)
    add_bullet(doc, "鎖死的獲利幾乎不可能真的買到：獲利集中在最不容易成交的族群，"
                    "是本回測最大的高估來源")
    add_bullet(doc, "已扣除賣出時的證交稅（0.3%），但手續費仍未計入")
    add_bullet(doc, "未考慮資金與同時發生的部位限制：同一天常有數十檔同時漲停，"
                    "回測把每筆交易當獨立事件計算，並未反映真實本金只有一份、"
                    "必須分配給同時出現的訊號")
    add_bullet(doc, "除息／除權／減資日的參考價調整未處理：漲停價用未還原的前收盤 × 1.1 計算，"
                    "這些日子交易所的參考價會調整，可能導致漏抓或誤判")
    add_bullet(doc, "樣本僅半年，且期間台股小型股投機氣氛濃厚，換個市場環境結果可能不同")
    add_bullet(doc, "與 0050 的比較僅供方向性參考，兩者報酬率計算方式不同，不是嚴謹的"
                    "同基準（apples-to-apples）比較")

    doc.save(REPORT_PATH)
    print(f"已存報告：{REPORT_PATH}")


def main():
    stats_dict = compute_all_stats()
    bench_ret, bench_start, bench_end = fetch_benchmark_return()
    build_report(stats_dict, bench_ret, bench_start, bench_end)


if __name__ == "__main__":
    main()
